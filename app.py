import io
import math
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# =========================================================
# Utilities (kept from your current code)
# =========================================================
def _step_status():
    """Derive step completion from session_state."""
    client_ok = st.session_state.get("client") is not None
    bench_ok = st.session_state.get("bench") is not None
    uploaded_ok = client_ok and bench_ok
    results_ok = bool(st.session_state.get("results_ready", False))
    return uploaded_ok, results_ok

def render_progress_header(current_step: int):
    """
    current_step: 1..4 corresponding to Upload/Configure/Results/Exports
    Renders a horizontal step indicator.
    """
    uploaded_ok, results_ok = _step_status()

    steps = [
        {"n": 1, "label": "Upload", "done": uploaded_ok},
        {"n": 2, "label": "Configure", "done": uploaded_ok},  # configure is available once uploads exist
        {"n": 3, "label": "Results", "done": results_ok},
        {"n": 4, "label": "Exports", "done": results_ok},
    ]

    cols = st.columns([1, 0.3, 1, 0.3, 1, 0.3, 1])
    col_idx = 0

    for i, s in enumerate(steps):
        # Step pill
        with cols[col_idx]:
            if s["done"]:
                badge = "✔"
            elif s["n"] == current_step:
                badge = "●"
            else:
                badge = "○"

            st.markdown(
                f"**{badge} {s['n']}) {s['label']}**",
                help="✔ complete • ● current • ○ not started"
            )

        # Arrow between steps
        if i < len(steps) - 1:
            with cols[col_idx + 1]:
                st.markdown("<div style='text-align:center; padding-top:6px'>→</div>", unsafe_allow_html=True)

        col_idx += 2

    st.markdown("---")

def next_step_hint(next_tab_label: str, *, disabled: bool = False, key: str):
    """
    Streamlit tabs cannot be programmatically selected. This adds a CTA and scroll-to-top.
    key MUST be unique per call site to avoid StreamlitDuplicateElementId.
    """
    cols = st.columns([3, 2])
    with cols[1]:
        if st.button("Next step →", disabled=disabled, use_container_width=True, key=key):
            st.markdown(
                "<script>window.scrollTo({top: 0, behavior: 'smooth'});</script>",
                unsafe_allow_html=True
            )
            st.info(f"Next: click **{next_tab_label}** at the top.")

def _safe_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in (" ", "_", "-") else "_" for c in name).strip().replace(" ", "_")

def _format_money(x):
    try:
        return f"${float(x):,.2f}"
    except Exception:
        return ""

def split_lane_detail(text: str):
    if not isinstance(text, str):
        return "", "", "", ""
    s = text.strip()
    if not s:
        return "", "", "", ""
    parts = s.split(" to ")
    if len(parts) != 2:
        return "", "", "", ""
    orig, dest = parts[0].strip(), parts[1].strip()

    def split_part(p):
        if "," not in p:
            return p.strip(), ""
        city, st_ = p.rsplit(",", 1)
        return city.strip(), st_.strip()

    oc, os_ = split_part(orig)
    dc, ds_ = split_part(dest)
    return oc, os_, dc, ds_

def ensure_origin_dest(df: pd.DataFrame) -> pd.DataFrame:
    needed = ["origin_city", "origin_state", "dest_city", "dest_state"]
    if all(c in df.columns for c in needed):
        return df

    if df.empty:
        for c in needed:
            if c not in df.columns:
                df[c] = ""
        return df

    if "lane_detail" in df.columns:
        src = df["lane_detail"]
    elif "Lane_Detail" in df.columns:
        src = df["Lane_Detail"]
    elif "_lane" in df.columns:
        src = df["_lane"]
    else:
        for c in needed:
            if c not in df.columns:
                df[c] = ""
        return df

    od = src.apply(
        lambda x: pd.Series(
            split_lane_detail(x),
            index=["origin_city", "origin_state", "dest_city", "dest_state"],
        )
    )
    for c in needed:
        if c not in df.columns:
            df[c] = od[c]
    return df

def _clean_cs(x):
    if pd.isna(x):
        return ""
    return str(x).strip().upper().replace(" ", "")

def norm_lane(x):
    if pd.isna(x):
        return ""
    return str(x).strip().upper()

def parse_percent_col(series: pd.Series) -> pd.Series:
    s = series.astype(str).str.strip().str.replace("%", "", regex=False)
    vals = pd.to_numeric(s, errors="coerce")
    max_val = vals.max(skipna=True)
    if max_val is not None and max_val <= 1.0:
        return vals
    return vals / 100.0

def read_any(upload, sheet=None):
    name = upload.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(upload)
    elif name.endswith((".xlsx", ".xlsm", ".xls")):
        return pd.read_excel(upload, sheet_name=sheet)
    elif name.endswith(".xlsb"):
        return pd.read_excel(upload, engine="pyxlsb", sheet_name=sheet)
    else:
        raise ValueError("Unsupported file type. Please upload CSV/XLSX/XLS/XLSB.")

@st.cache_data
def read_any_cached(upload, sheet=None):
    if upload is None:
        return None
    return read_any(upload, sheet=sheet)

@st.cache_data
def infer_sheets_cached(upload):
    if upload is None:
        return []
    name = upload.name.lower()
    if name.endswith((".xlsx", ".xlsm", ".xls", ".xlsb")):
        try:
            xl = pd.ExcelFile(upload)
            return xl.sheet_names
        except Exception:
            return []
    return []

def guess_col(columns, candidates):
    if not columns:
        return 0
    lowered = [str(c).lower() for c in columns]
    for cand in candidates:
        for i, col in enumerate(lowered):
            if cand in col:
                return i
    return 0

# =========================================================
# Letter builders (kept from your current code)
# =========================================================

def build_letter_docx(
    carrier: str,
    df_carrier: pd.DataFrame,
    sender_company: str,
    sender_name: str,
    sender_title: str,
    reply_by: str,
    body_template: str,
) -> bytes:
    df_carrier = ensure_origin_dest(df_carrier.copy())
    doc = Document()

    title = doc.add_paragraph(f"Rate Review Request — {carrier}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.bold = True

    num_lanes = len(df_carrier)
    if "delta_pct" in df_carrier.columns:
        avg_over_pct = df_carrier["delta_pct"].mean(skipna=True)
        avg_over_pct_text = f"{avg_over_pct:.0f}%" if pd.notna(avg_over_pct) else "N/A"
    else:
        avg_over_pct_text = "N/A"

    body_text = body_template.format(
        num_lanes=num_lanes,
        avg_over_pct=avg_over_pct_text,
        reply_by=reply_by,
    )
    doc.add_paragraph(body_text)

    total_delta = df_carrier["delta"].sum(skipna=True) if "delta" in df_carrier.columns else 0
    doc.add_paragraph(
        f"Lanes to negotiate: {num_lanes} • Current total variance vs. benchmark: {_format_money(total_delta)}"
    )

    lane_group = (
        df_carrier
        .groupby(["origin_city", "origin_state", "dest_city", "dest_state"], dropna=False)
        .agg(
            Frequency=("_lane", "size"),
            Company_Cost=("company_cost", "mean"),
            Benchmark_Cost=("benchmark_cost", "mean"),
            Over_Pct=("delta_pct", "mean"),
        )
        .reset_index()
    )

    pretty = lane_group[
        ["origin_city", "origin_state", "dest_city", "dest_state", "Frequency", "Company_Cost", "Benchmark_Cost", "Over_Pct"]
    ].copy()

    pretty.columns = [
        "Origin City",
        "Origin State",
        "Dest City",
        "Dest State",
        "Frequency",
        "Company Cost",
        "Benchmark Cost",
        "% over Benchmark",
    ]

    table = doc.add_table(rows=1, cols=len(pretty.columns))
    hdr = table.rows[0].cells
    for i, c in enumerate(pretty.columns):
        p = hdr[i].paragraphs[0]
        run = p.add_run(c)
        run.bold = True
        run.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _, row in pretty.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Origin City"])
        cells[1].text = str(row["Origin State"])
        cells[2].text = str(row["Dest City"])
        cells[3].text = str(row["Dest State"])
        cells[4].text = str(int(row["Frequency"])) if not pd.isna(row["Frequency"]) else ""
        cells[5].text = _format_money(row["Company Cost"])
        cells[6].text = _format_money(row["Benchmark Cost"])
        cells[7].text = f"{row['% over Benchmark']:.0f}%" if pd.notna(row["% over Benchmark"]) else ""

    note = doc.add_paragraph(
        " Benchmark cost reflects market-based rates for comparable lanes and serves as a reference point for rate review."
    )
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(12)
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)

    doc.add_paragraph("")
    doc.add_paragraph("We appreciate your partnership and look forward to your response.")
    doc.add_paragraph(f"\n{sender_name}\n{sender_title}\n{sender_company}")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def build_letters_zip(df_all, include_privfleet: bool, sender_company: str, sender_name: str,
                      sender_title: str, reply_by: str, body_template: str) -> bytes:
    if not include_privfleet and "carrier_name" in df_all.columns:
        df_all = df_all[df_all["carrier_name"].astype(str).str.upper() != "GREIF PRIVATE FLEET"]

    df_neg = df_all[df_all["action"] == "NEGOTIATE"].copy() if "action" in df_all.columns else df_all.iloc[0:0].copy()
    if df_neg.empty:
        buff = BytesIO()
        with ZipFile(buff, "w", ZIP_DEFLATED):
            pass
        buff.seek(0)
        return buff.read()

    buff = BytesIO()
    with ZipFile(buff, "w", ZIP_DEFLATED) as zf:
        for carrier, dfc in df_neg.groupby("carrier_name", dropna=False):
            carrier_name = "UNKNOWN" if pd.isna(carrier) or str(carrier).strip() == "" else str(carrier)
            dfc = dfc.sort_values("delta", ascending=False, na_position="last")

            doc_bytes = build_letter_docx(
                carrier=carrier_name,
                df_carrier=dfc,
                sender_company=sender_company,
                sender_name=sender_name,
                sender_title=sender_title,
                reply_by=reply_by,
                body_template=body_template,
            )
            fname = f"Negotiation_{_safe_filename(carrier_name)}.docx"
            zf.writestr(fname, doc_bytes)

    buff.seek(0)
    return buff.read()

# =========================================================
# App config
# =========================================================

st.set_page_config(page_title="Freight Lane Comparison", layout="wide")
st.markdown(
    """
    <style>
    .flo-header {
        display: flex;
        align-items: center;
        gap: 16px;            /* controls logo ↔ text distance */
        margin-top: 8px;
        margin-bottom: 8px;
    }
    .flo-header img {
        width: 200px;        /* logo size */
        height: auto;
    }
    .flo-title {
        display: flex;
        flex-direction: column;
        justify-content: center;
    }
    .flo-title h1 {
        margin: 0;
        font-size: 34px;
        line-height: 1.1;
    }
    .flo-title p {
        margin: 4px 0 0 0;
        font-size: 15px;
        color: #6b7280;
    }
    </style>

    <div class="flo-header">
        <img src="assets/flo_logo.png" alt="FLO.ai logo">
        <div class="flo-title">
            <h1>FLO.ai</h1>
            <p>Freight Lane Optimization</p>
        </div>
    </div>
    <hr>
    """,
    unsafe_allow_html=True
)

# =========================================================
# Top-of-page: What this tool does + toggleable How-to
# =========================================================

st.markdown("### What this tool does")

st.markdown(
    """
This app compares your **company freight costs** to **benchmark rates**, then classifies each lane into:

- **RFP lanes**: lanes that will be included in a broader bid event
- **Negotiation lanes**: *non-vanilla* lanes or lanes excluded from RFP for specific reasons that  
  will be handled via targeted vendor letters. Letters can also be used for monthly rate review,  
  e.g., month-to-month negotiations with carriers to monitor variance to benchmark
- **Excluded lanes**: filtered out due to location, mode, or carrier exclusions

From this, you can:

- Download a **full comparison workbook**
- Generate an **RFP template** (Overview, Timeline, Bid Lanes, Locations)
- Generate **RFP letters** to carriers participating in the bid
- Generate **negotiation letters** for *non-vanilla* lanes flagged for direct rate review

**Note:** When the program is running or loading, the screen will temporarily *gray out*.  
Please **do not refresh the page or change inputs** while the page is loading.
"""
)

# ---- Toggle state for how-to ----
if "show_howto" not in st.session_state:
    st.session_state["show_howto"] = False

# Toggle button
toggle_label = "Show: How to use this tool" if not st.session_state["show_howto"] else "Hide: How to use this tool"
if st.button(toggle_label, key="toggle_howto_top"):
    st.session_state["show_howto"] = not st.session_state["show_howto"]
    st.rerun()

# Conditionally render how-to content
if st.session_state["show_howto"]:
    st.markdown("---")
    st.markdown("### How to use this tool")

    st.markdown(
        "**Please note:** When the app is running, the page will temporarily gray out. "
        "Do not refresh or change inputs while processing.\n\n"

        "**Step 1: Upload data**\n"
        "1. Upload **Company** and **Benchmark** files\n"
        "2. Map the correct columns (lane, cost, carrier, mode)\n"
        "3. Add any **location, carrier, or mode exclusions**\n\n"

        "**Step 2: Run comparison**\n"
        "- Click **Run comparison** to:\n"
        "  - Match lanes to benchmark\n"
        "  - Compute $ and % deltas\n"
        "  - Classify lanes into RFP vs Negotiation vs Excluded\n\n"

        "**When to generate RFP Template + RFP Letters**\n"
        "- Use when running a **formal bid event** across multiple lanes/carriers\n"
        "- Recommended no more than **twice per year**\n"
        "- RFP Template includes Overview, Timeline, Bid Lanes, Location List\n"
        "- RFP Letters summarize above-benchmark lanes and highlight top 10% variances\n\n"

        "**When to generate Negotiation Letters**\n"
        "- Use when you **do not** want a full RFP\n"
        "- Target specific lanes for **direct rate negotiation**\n"
        "- Useful as a **monthly compliance check** for above-benchmark lanes"
    )

    st.caption(
        "Tip: Upload → Configure → Run comparison → Review Results → Generate Exports"
    )

st.markdown("---")

FIXED_EXCLUDE_LOCATIONS = [
    "GLADSTONEVA", "FITCHBURGMA", "MORGAN HILLCA", "MASSILLONOH", "MERCEDCA",
    "LOUISVILLEKY", "MASONMI", "GREENSBORONC", "CONCORDNC", "PALMYRAPA", "DALLASTX"
]

# =========================================================
# Session state (IMPORTANT: never None for results tables)
# =========================================================

if "results_ready" not in st.session_state:
    st.session_state["results_ready"] = False
if "out" not in st.session_state:
    st.session_state["out"] = pd.DataFrame()
if "gpf_export" not in st.session_state:
    st.session_state["gpf_export"] = pd.DataFrame()
if "summary_df" not in st.session_state:
    st.session_state["summary_df"] = pd.DataFrame()
if "excluded_summary_df" not in st.session_state:
    st.session_state["excluded_summary_df"] = pd.DataFrame(columns=["Excluded_Location", "Count"])
if "excluded_detail_df" not in st.session_state:
    st.session_state["excluded_detail_df"] = pd.DataFrame(columns=["_lane", "carrier_name", "excluded_matches"])

# =========================================================
# Tabs (1 + 10)
# =========================================================
st.markdown('<div id="top"></div>', unsafe_allow_html=True)

tab_upload, tab_config, tab_results, tab_exports = st.tabs(
    ["1) Upload", "2) Configure", "3) Results", "4) Exports"]
)

# =========================================================
# 1) Upload
# =========================================================

with tab_upload:
    render_progress_header(current_step=1)
    colL, colR = st.columns(2)

    with colL:
        st.subheader("Company file")
        client_file = st.file_uploader(
            "Upload Company data (CSV/XLSX/XLS/XLSB)",
            type=["csv", "xlsx", "xls", "xlsb"],
            key="client",
        )
        client_sheets = infer_sheets_cached(client_file)
        client_sheet = st.selectbox(
            "Company sheet (optional)",
            options=["<first sheet>"] + client_sheets if client_sheets else ["<first sheet>"],
        )

        st.caption("Optional: Use this template if you need a structured company file.")
        template_cols = [
            "Origin City", "Origin State", "Dest City", "Dest State",
            "Total Base Charges", "Carrier Name", "Carrier Mode",
        ]
        template_df = pd.DataFrame(columns=template_cols)
        tmpl_buf = io.BytesIO()
        with pd.ExcelWriter(tmpl_buf, engine="openpyxl") as writer:
            template_df.to_excel(writer, index=False, sheet_name="Company Data Template")
        tmpl_buf.seek(0)
        st.download_button(
            label="⬇️ Download Company Data Template (Excel)",
            data=tmpl_buf,
            file_name="company_data_template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    with colR:
        st.subheader("Benchmark file")
        bench_file = st.file_uploader(
            "Upload Benchmark (CSV/XLSX/XLS/XLSB)",
            type=["csv", "xlsx", "xls", "xlsb"],
            key="bench",
        )
        bench_sheets = infer_sheets_cached(bench_file)
        bench_sheet = st.selectbox(
            "Benchmark sheet (optional)",
            options=["<first sheet>"] + bench_sheets if bench_sheets else ["<first sheet>"],
        )

    # Previews (for automapping dropdown defaults)
    client_mode_columns = ["<None>"]
    bench_mode_columns = ["<None>"]
    df_client_preview = pd.DataFrame()
    df_bench_preview = pd.DataFrame()

    if client_file is not None:
        try:
            sheet_c_preview = None if client_sheet == "<first sheet>" else client_sheet
            df_client_preview = read_any_cached(client_file, sheet_c_preview)
            client_mode_columns = ["<None>"] + list(df_client_preview.columns)
            st.success(f"Company file loaded: {df_client_preview.shape[0]:,} rows, {df_client_preview.shape[1]:,} cols")
        except Exception as e:
            st.error(f"Could not read Company file: {e}")

    if bench_file is not None:
        try:
            sheet_b_preview = None if bench_sheet == "<first sheet>" else bench_sheet
            df_bench_preview = read_any_cached(bench_file, sheet_b_preview)
            bench_mode_columns = ["<None>"] + list(df_bench_preview.columns)
            st.success(f"Benchmark file loaded: {df_bench_preview.shape[0]:,} rows, {df_bench_preview.shape[1]:,} cols")
        except Exception as e:
            st.error(f"Could not read Benchmark file: {e}")
    client_file = st.session_state.get("client")
    bench_file = st.session_state.get("bench")
    next_step_hint("2) Configure", disabled=(client_file is None or bench_file is None), key="next_upload")

# =========================================================
# 2) Configure (2 + 3 + 7)
# =========================================================

with tab_config:
    render_progress_header(current_step=2)
    st.subheader("Configure comparison")

    if (st.session_state.get("client") is None) and (st.session_state.get("bench") is None):
        st.info("Upload files in the Upload tab first (recommended), then configure mappings here.")

    with st.form("config_form", clear_on_submit=False):
        st.subheader("Required mappings")

        c1, c2, c3, c4 = st.columns(4)

        with c1:
            client_lane_col = st.selectbox(
                "Company lane column (key)",
                options=client_mode_columns,
                index=guess_col(client_mode_columns, ["_lane", "lane", "od", "origin"]),
            )
            company_cost_col = st.selectbox(
                "Company cost column",
                options=client_mode_columns,
                index=guess_col(client_mode_columns, ["cost", "charge", "linehaul", "rate", "amount", "base"]),
            )
            client_carrier_col = st.selectbox(
                "Company carrier column",
                options=client_mode_columns,
                index=guess_col(client_mode_columns, ["carrier", "vendor", "scac"]),
            )
            lane_detail_col = st.text_input(
                "Company lane detail column (optional)",
                value="Lane_Detail",
            )

        with c2:
            bench_lane_col = st.selectbox(
                "Benchmark lane column (key)",
                options=bench_mode_columns,
                index=guess_col(bench_mode_columns, ["_lane", "lane", "od", "origin"]),
            )
            bench_cost_col = st.selectbox(
                "Benchmark cost column",
                options=bench_mode_columns,
                index=guess_col(bench_mode_columns, ["cost", "charge", "linehaul", "rate", "amount", "base"]),
            )
            bench_agg = st.selectbox("Benchmark duplicate lanes aggregation", options=["mean", "median"], index=0)

        with c3:
            st.markdown("**Optional: fuel**")
            company_fuel_col = st.selectbox(
                "Company fuel surcharge $ column (optional)",
                options=client_mode_columns,
                index=guess_col(client_mode_columns, ["fuel"]),
            )
            bench_fuel_col = st.selectbox(
                "Benchmark fuel surcharge % column (optional)",
                options=bench_mode_columns,
                index=guess_col(bench_mode_columns, ["fuel"]),
            )

        with c4:
            st.markdown("**Optional: mode**")
            mode_col_client = st.selectbox(
                "Company mode column (TL/LTL)",
                options=client_mode_columns,
                index=guess_col(client_mode_columns, ["mode", "tl", "ltl"]),
            )
            bench_mode_col = st.selectbox(
                "Benchmark mode column (TL/LTL)",
                options=bench_mode_columns,
                index=guess_col(bench_mode_columns, ["mode", "tl", "ltl"]),
            )

        st.markdown("---")
        st.subheader("Advanced: Exclusions")

        ex1, ex2 = st.columns(2)
        with ex1:
            apply_fixed_exclusions = st.checkbox("Apply default location exclusions", value=True)
            if apply_fixed_exclusions:
                st.caption("Default excluded locations:")
                st.code(", ".join(FIXED_EXCLUDE_LOCATIONS), language=None)

            extra_locations = st.text_area(
                "Extra locations to exclude (comma-separated)",
                placeholder="ATLANTAGA, CHICAGOIL, BOSTONMA",
            )

        with ex2:
            extra_carriers = st.text_area(
                "Carriers to exclude (comma-separated)",
                placeholder="CARRIER A, CARRIER B",
            )

        st.markdown("---")
        st.subheader("Mode handling")

        match_by_mode = st.checkbox(
            "Match by mode when both mode columns are mapped",
            value=True,
        )

        exclude_modes_raw = st.text_input(
            "Modes to exclude (comma-separated, case-insensitive)",
            value="LTL",
        )

        st.markdown("---")
        st.subheader("RFP overrides")
        letter_override_raw = st.text_area(
            "Lane keys to treat as Vendor Letters instead of RFP (comma or newline separated)",
            placeholder="ATLANTAGADALLASTX, CHICAGOILHOUSTONTX",
        )

        submitted = st.form_submit_button("Run comparison")

    # Normalize override list outside form widget creation
    override_letter_lanes = set()
    if str(letter_override_raw).strip():
        pieces = [c.strip() for c in str(letter_override_raw).replace("\n", ",").split(",") if c.strip()]
        override_letter_lanes = {p.upper() for p in pieces}
    
    # Store overrides so Results/Exports can reliably access them
    st.session_state["override_letter_lanes"] = override_letter_lanes
    
    next_step_hint("3) Results", disabled=(not st.session_state.get("results_ready", False)), key="next_config")

# =========================================================
# Run comparison (2 + 6)
# =========================================================

if submitted:
    with st.status("Running comparison…", expanded=True) as status:
        st.write("Validating inputs")

        # Pull the file objects from Streamlit widget state
        client_file = st.session_state.get("client")
        bench_file = st.session_state.get("bench")

        if client_file is None or bench_file is None:
            status.update(label="Missing inputs", state="error")
            st.error("Please upload both Company and Benchmark files in the Upload tab.")
            st.stop()

        sheet_c = None if client_sheet == "<first sheet>" else client_sheet
        sheet_b = None if bench_sheet == "<first sheet>" else bench_sheet

        st.write("Reading files")
        df_client = read_any_cached(client_file, sheet_c)
        df_bench = read_any_cached(bench_file, sheet_b)

        if df_client is None or df_bench is None:
            status.update(label="Failed to read files", state="error")
            st.error("Unable to read one or both files.")
            st.stop()

        st.write("Mode matching setup")
        client_has_mode = (mode_col_client not in ("<None>", None, "") and mode_col_client in df_client.columns)
        bench_has_mode = (bench_mode_col not in ("<None>", None, "") and bench_mode_col in df_bench.columns)
        use_mode_matching = bool(match_by_mode and client_has_mode and bench_has_mode)
        # ---------------------------------------------------------
        # SAFETY RULE:
        # If benchmark has no mode column, exclude LTL from company
        # to avoid false TL vs LTL comparisons
        # ---------------------------------------------------------
        auto_exclude_ltl_due_to_benchmark = not bench_has_mode

        st.write("Normalizing lanes")
        df_client = df_client.copy()
        df_bench = df_bench.copy()

        # Ensure parsed origin/dest exists for later exports/letters
        df_client = ensure_origin_dest(df_client)
        df_bench = ensure_origin_dest(df_bench)

        # Validate required columns
        missing_client = [c for c in [client_lane_col, company_cost_col, client_carrier_col] if c not in df_client.columns]
        missing_bench = [c for c in [bench_lane_col, bench_cost_col] if c not in df_bench.columns]
        if missing_client:
            status.update(label="Missing company columns", state="error")
            st.error(f"Company file missing columns: {missing_client}")
            st.stop()
        if missing_bench:
            status.update(label="Missing benchmark columns", state="error")
            st.error(f"Benchmark file missing columns: {missing_bench}")
            st.stop()

        df_client["_lane"] = df_client[client_lane_col].map(norm_lane)
        df_bench["_lane"] = df_bench[bench_lane_col].map(norm_lane)

        def norm_mode(series):
            return (
                series.astype(str)
                .str.strip()
                .str.upper()
                .replace({
                    "TRUCKLOAD": "TL",
                    "TL": "TL",
                    "LTL": "LTL",
                    "LESS THAN TRUCKLOAD": "LTL",
                })
            )

        # Always normalize company mode if column exists
        if client_has_mode:
            df_client["_mode"] = norm_mode(df_client[mode_col_client])
        else:
            df_client["_mode"] = "UNKNOWN"
        
        # Only normalize benchmark mode if available
        if bench_has_mode:
            df_bench["_mode"] = norm_mode(df_bench[bench_mode_col])
        else:
            df_bench["_mode"] = "UNKNOWN"

        st.write("Computing costs")
        df_client["company_linehaul"] = pd.to_numeric(df_client[company_cost_col], errors="coerce")

        company_has_fuel = (company_fuel_col not in (None, "", "<None>") and company_fuel_col in df_client.columns)
        if company_has_fuel:
            df_client["company_fuel_cost"] = pd.to_numeric(df_client[company_fuel_col], errors="coerce")
        else:
            df_client["company_fuel_cost"] = 0.0

        df_client["company_cost"] = df_client["company_linehaul"] + df_client["company_fuel_cost"]

        df_bench["benchmark_linehaul"] = pd.to_numeric(df_bench[bench_cost_col], errors="coerce")

        if bench_fuel_col != "<None>" and bench_fuel_col in df_bench.columns:
            bench_fuel_pct = parse_percent_col(df_bench[bench_fuel_col])
            df_bench["benchmark_fuel_cost"] = df_bench["benchmark_linehaul"] * bench_fuel_pct
        else:
            df_bench["benchmark_fuel_cost"] = 0.0

        df_bench["benchmark_cost"] = df_bench["benchmark_linehaul"] + df_bench["benchmark_fuel_cost"]

        st.write("Applying exclusions (locations/carriers/modes)")

        client_keep = df_client[[
            "_lane", client_carrier_col, "company_linehaul", "company_fuel_cost", "company_cost",
            "_mode", "origin_city", "origin_state", "dest_city", "dest_state"
        ]].rename(columns={client_carrier_col: "carrier_name", "_mode": "mode"})

        if lane_detail_col in df_client.columns:
            client_keep["lane_detail"] = df_client[lane_detail_col]

        # Location exclusions
        exclude_locations = FIXED_EXCLUDE_LOCATIONS.copy() if apply_fixed_exclusions else []
        if extra_locations:
            exclude_locations.extend([x.strip().upper() for x in extra_locations.split(",") if x.strip()])

        client_keep["lane_key_upper"] = client_keep["_lane"].astype(str).str.upper()

        def match_locs(text: str) -> list:
            return [loc for loc in exclude_locations if loc in text]

        client_keep["excluded_matches"] = client_keep["lane_key_upper"].apply(match_locs)
        mask_loc_excl = client_keep["excluded_matches"].apply(lambda v: len(v) if isinstance(v, list) else 0).gt(0)

        excluded_detail_df = client_keep.loc[mask_loc_excl, ["_lane", "carrier_name", "excluded_matches"]].copy()
        excluded_exploded = excluded_detail_df.explode("excluded_matches")
        if len(excluded_exploded):
            excluded_summary_df = (
                excluded_exploded["excluded_matches"]
                .value_counts(dropna=False)
                .rename_axis("Excluded_Location")
                .reset_index(name="Count")
            )
        else:
            excluded_summary_df = pd.DataFrame(columns=["Excluded_Location", "Count"])

        client_keep = client_keep.loc[~mask_loc_excl].drop(columns=["excluded_matches"], errors="ignore")

        # Carrier exclusions
        if extra_carriers:
            carriers_list = [c.strip().upper() for c in extra_carriers.split(",") if c.strip()]
            client_keep = client_keep[~client_keep["carrier_name"].astype(str).str.upper().isin(carriers_list)]

        # Mode exclusions (explicit + automatic)
        exclude_modes = {m.strip().upper() for m in str(exclude_modes_raw).split(",") if m.strip()}
        
        # Auto-exclude LTL if benchmark has no mode info
        if auto_exclude_ltl_due_to_benchmark:
            exclude_modes.add("LTL")
        
        if exclude_modes:
            if "mode" in client_keep.columns:
                client_keep = client_keep[
                    ~client_keep["mode"].astype(str).str.upper().isin(exclude_modes)
                ]
        # ---------------------------------------------------------
        # User-facing explanation for automatic LTL exclusion
        # ---------------------------------------------------------
        if auto_exclude_ltl_due_to_benchmark:
            st.info(
                "ℹ️ **LTL lanes were automatically excluded** because the benchmark file "
                "does not contain a usable mode column. "
                "This prevents TL/LTL mismatches and false savings comparisons."
            )

        st.write("Aggregating benchmark and merging")
        bench_keep = df_bench[["_lane", "_mode", "benchmark_linehaul", "benchmark_fuel_cost", "benchmark_cost"]].rename(columns={"_mode": "mode"})

        group_cols = ["_lane"] + (["mode"] if use_mode_matching else [])
        value_cols = ["benchmark_linehaul", "benchmark_fuel_cost", "benchmark_cost"]
        agg_func = "median" if bench_agg == "median" else "mean"

        bench_agg_df = bench_keep.groupby(group_cols, as_index=False, dropna=False)[value_cols].agg(agg_func)

        if use_mode_matching:
            merged = client_keep.merge(bench_agg_df, how="left", on=["_lane", "mode"])
        else:
            merged = client_keep.merge(bench_agg_df, how="left", on=["_lane"])

        merged["has_benchmark"] = merged["benchmark_cost"].notna() & (merged["benchmark_cost"] > 0)
        merged["benchmark_linehaul"] = merged["benchmark_linehaul"].fillna(0.0)
        merged["benchmark_fuel_cost"] = merged["benchmark_fuel_cost"].fillna(0.0)
        merged["benchmark_cost"] = merged["benchmark_cost"].fillna(0.0)

        merged["delta_linehaul"] = merged["company_linehaul"] - merged["benchmark_linehaul"]
        merged["delta_fuel"] = merged["company_fuel_cost"] - merged["benchmark_fuel_cost"]
        merged["delta"] = merged["company_cost"] - merged["benchmark_cost"]

        merged["delta_pct"] = pd.NA
        pct_mask = merged["has_benchmark"] & (merged["benchmark_cost"] != 0)
        merged.loc[pct_mask, "delta_pct"] = (merged.loc[pct_mask, "delta"] / merged.loc[pct_mask, "benchmark_cost"] * 100.0)

        def decide_action(row):
            if not row["has_benchmark"]:
                return "None"
            if pd.isna(row["delta"]) or row["delta"] <= 0:
                return "None"
            return "NEGOTIATE"

        merged["action"] = merged.apply(decide_action, axis=1)

        out = merged[[
            "_lane", "carrier_name", "mode",
            "company_linehaul", "company_fuel_cost", "company_cost",
            "benchmark_linehaul", "benchmark_fuel_cost", "benchmark_cost",
            "delta_linehaul", "delta_fuel", "delta", "delta_pct",
            "action", "origin_city", "origin_state", "dest_city", "dest_state"
        ]].sort_values("delta", ascending=False, na_position="last")

        st.write("Preparing private fleet view")
        gpf_export = client_keep[client_keep["carrier_name"].astype(str).str.upper() == "GREIF PRIVATE FLEET"].copy()
        if gpf_export.empty:
            gpf_export = pd.DataFrame(columns=["_lane", "carrier_name", "company_cost", "benchmark_cost", "delta", "action"])
        else:
            # Merge private fleet to benchmark agg (use same matching basis)
            if use_mode_matching:
                gpf_m = gpf_export.merge(bench_agg_df, how="left", on=["_lane", "mode"])
            else:
                gpf_m = gpf_export.merge(bench_agg_df, how="left", on=["_lane"])
            gpf_m["benchmark_cost"] = gpf_m["benchmark_cost"].fillna(0.0)
            gpf_m["delta"] = gpf_m["company_cost"] - gpf_m["benchmark_cost"]
            gpf_m["action"] = gpf_m["delta"].apply(lambda d: "NEGOTIATE" if pd.notna(d) and d > 0 else "None")
            gpf_export = gpf_m[["_lane", "carrier_name", "company_cost", "benchmark_cost", "delta", "action"]]

        gpf_export = ensure_origin_dest(gpf_export)

        st.write("Building summary")
        out_no_pf = out[out["carrier_name"].astype(str).str.upper() != "GREIF PRIVATE FLEET"].copy()

        neg_mask_out = out_no_pf["action"] == "NEGOTIATE"
        overall_count = int(neg_mask_out.sum())
        overall_total = float(out_no_pf.loc[neg_mask_out, "delta"].sum(skipna=True))

        gpf_negotiate_count = int((gpf_export["action"] == "NEGOTIATE").sum()) if not gpf_export.empty else 0
        gpf_total_delta = float(gpf_export.loc[gpf_export["action"] == "NEGOTIATE", "delta"].sum(skipna=True)) if not gpf_export.empty else 0.0
        gpf_count = int(gpf_export.shape[0]) if isinstance(gpf_export, pd.DataFrame) else 0

        summary_df = pd.DataFrame([
            {
                "Segment": "OVERALL (Non-Private Fleet)",
                "Negotiate_Lanes": overall_count,
                "Total_Delta": overall_total,
                "Summary_Text": f"SUMMARY (OVERALL): {overall_count} lanes marked as NEGOTIATE with total delta ${overall_total:,.2f}."
            },
            {
                "Segment": "PRIVATE FLEET",
                "Negotiate_Lanes": gpf_negotiate_count,
                "Total_Delta": gpf_total_delta,
                "Summary_Text": f"SUMMARY: {gpf_count} lanes total; NEGOTIATE lanes delta ${gpf_total_delta:,.2f}."
            }
        ])

        # Persist results
        st.session_state["out"] = out_no_pf
        st.session_state["gpf_export"] = gpf_export
        st.session_state["summary_df"] = summary_df
        st.session_state["excluded_summary_df"] = excluded_summary_df
        st.session_state["excluded_detail_df"] = excluded_detail_df
        st.session_state["results_ready"] = True

        status.update(label="Comparison complete", state="complete")

# =========================================================
# 3) Results (8) – fully gated, no global unguarded access
# =========================================================

with tab_results:
    render_progress_header(current_step=3)
    st.header("📌 Decision Summary")

    if not st.session_state["results_ready"]:
        st.info("Run the comparison in the Configure tab to see results.")
    else:
        out = st.session_state["out"]
        gpf_export = st.session_state["gpf_export"]
        summary_df = st.session_state["summary_df"]
        excluded_summary_df = st.session_state["excluded_summary_df"]
        excluded_detail_df = st.session_state["excluded_detail_df"]
        override_letter_lanes = st.session_state.get("override_letter_lanes", set())

        # KPIs (8)
        carriers_impacted = int(out.loc[out["action"] == "NEGOTIATE", "carrier_name"].nunique()) if "carrier_name" in out.columns else 0
        total_lanes = int(out.shape[0]) if isinstance(out, pd.DataFrame) else 0
        negotiate_lanes = int((out["action"] == "NEGOTIATE").sum()) if total_lanes else 0
        total_savings = float(out.loc[out["action"] == "NEGOTIATE", "delta"].clip(lower=0).sum(skipna=True)) if total_lanes else 0.0
        match_rate = float(out["benchmark_cost"].notna().mean()) if total_lanes and "benchmark_cost" in out.columns else 0.0
        excluded_lanes = int(excluded_detail_df.shape[0]) if isinstance(excluded_detail_df, pd.DataFrame) else 0

        k1, k2, k3, k4, k5 = st.columns(5)
        k1.metric("Lanes analyzed", f"{total_lanes:,}")
        k2.metric("Above benchmark", f"{negotiate_lanes:,}")
        k3.metric("Savings potential (Δ>0)", f"${total_savings:,.0f}")
        k4.metric("Carriers impacted", f"{carriers_impacted:,}")
        k5.metric("Excluded lanes", f"{excluded_lanes:,}")

        st.markdown("---")
        # =========================
        # Decision Summary (Option 7)
        # =========================
        # Reuse what you already computed above:
        # total_lanes, negotiate_lanes, total_savings, match_rate, excluded_lanes
        
        # Derive carrier impact + RFP/Letter counts (lightweight, no extra heavy compute)
        out_tmp = out.copy()
        out_tmp["lane_key_norm"] = out_tmp["_lane"].astype(str).str.strip().str.upper()
        
        def _classify_for_summary(row):
            if row.get("action") != "NEGOTIATE":
                return "None"
            if pd.isna(row.get("benchmark_cost")) or row.get("benchmark_cost", 0) <= 0:
                return "None"
            if row["lane_key_norm"] in override_letter_lanes:
                return "Letter"
            return "RFP"
        
        out_tmp["lane_treatment"] = out_tmp.apply(_classify_for_summary, axis=1)
        
        rfp_count = int((out_tmp["lane_treatment"] == "RFP").sum())
        letter_count = int((out_tmp["lane_treatment"] == "Letter").sum())
        carriers_impacted = int(out_tmp.loc[out_tmp["action"] == "NEGOTIATE", "carrier_name"].nunique()) if "carrier_name" in out_tmp.columns else 0
        
        # A concise “what next” recommendation
        if negotiate_lanes == 0:
            next_reco = "No lanes are currently above benchmark under the selected settings. Consider relaxing exclusions or validating mappings."
        elif rfp_count > 0 and letter_count > 0:
            next_reco = "Review RFP lanes and Letter lanes, then proceed to Exports to build the workbook and communications."
        elif rfp_count > 0:
            next_reco = "Review RFP lanes, then proceed to Exports to build the RFP template and carrier letters."
        else:
            next_reco = "Review Letter lanes, then proceed to Exports to generate negotiation letters."
                
        st.caption(
            f"RFP candidates: {rfp_count:,} • Letter candidates: {letter_count:,} • Excluded lanes: {excluded_lanes:,}"
        )
        
        st.info(f"**Recommended next step:** {next_reco}")
        
        st.markdown("---")

        # Treatment classification (kept from your logic)
        out = out.copy()
        out["lane_key_norm"] = out["_lane"].astype(str).str.strip().str.upper()

        def classify_row(row):
            if row["action"] != "NEGOTIATE":
                return "None"
            if pd.isna(row["benchmark_cost"]) or row["benchmark_cost"] <= 0:
                return "None"
            if row["lane_key_norm"] in override_letter_lanes:
                return "Letter"
            return "RFP"

        out["lane_treatment"] = out.apply(classify_row, axis=1)

        rfp_df = out[out["lane_treatment"] == "RFP"].copy()
        letter_df = out[out["lane_treatment"] == "Letter"].copy()
        no_action_df = out[out["lane_treatment"] == "None"].copy()

        # Output tabs
        tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
            "🧾 Summary",
            "📦 RFP Candidates",
            "📊 Letter Candidates",
            "🚛 Private Fleet",
            "🚫 Excluded (Summary)",
            "🚫 Excluded (Detail)"
        ])

        with tab1:
            neg_mask = out["action"] == "NEGOTIATE"
            overall_total = float(out.loc[neg_mask, "delta"].sum(skipna=True))
            linehaul_savings = float(out.loc[neg_mask, "delta_linehaul"].sum(skipna=True))
            fuel_savings = float(out.loc[neg_mask, "delta_fuel"].sum(skipna=True))

            st.markdown("#### Savings scenarios")
            c1, c2, c3 = st.columns(3)
            c1.metric("Scenario 1 (Linehaul only)", f"${linehaul_savings:,.2f}")
            c2.metric("Scenario 2 (Linehaul + fuel)", f"${overall_total:,.2f}")
            c3.metric("Fuel component", f"${fuel_savings:,.2f}")

            st.dataframe(summary_df, use_container_width=True)

        with tab2:
            if rfp_df.empty:
                st.info("No lanes flagged for RFP.")
            else:
                st.dataframe(rfp_df.head(1000), use_container_width=True)

        with tab3:
            if letter_df.empty:
                st.info("No lanes flagged for Letters.")
            else:
                st.dataframe(letter_df.head(1000), use_container_width=True)

        with tab4:
            if gpf_export.empty:
                st.info("No PRIVATE FLEET lanes found after exclusions.")
            else:
                st.dataframe(gpf_export, use_container_width=True)

        with tab5:
            st.dataframe(excluded_summary_df, use_container_width=True)

        with tab6:
            st.dataframe(excluded_detail_df, use_container_width=True)
    
    next_step_hint("4) Exports", disabled=(not st.session_state.get("results_ready", False)), key="next_results")

# =========================================================
# 4) Exports (10) – fully gated; nothing runs without results
# =========================================================

with tab_exports:
    render_progress_header(current_step=4)
    st.subheader("Exports")

    if not st.session_state["results_ready"]:
        st.info("Run the comparison first to enable exports.")
    else:
        out = st.session_state["out"]
        gpf_export = st.session_state["gpf_export"]
        summary_df = st.session_state["summary_df"]
        excluded_summary_df = st.session_state["excluded_summary_df"]
        excluded_detail_df = st.session_state["excluded_detail_df"]

        # Rebuild treatments in exports tab (consistent with Results tab)
        out = out.copy()
        out["lane_key_norm"] = out["_lane"].astype(str).str.strip().str.upper()

        def classify_row(row):
            if row["action"] != "NEGOTIATE":
                return "None"
            if pd.isna(row["benchmark_cost"]) or row["benchmark_cost"] <= 0:
                return "None"
            if row["lane_key_norm"] in override_letter_lanes:
                return "Letter"
            return "RFP"

        out["lane_treatment"] = out.apply(classify_row, axis=1)
        rfp_df = out[out["lane_treatment"] == "RFP"].copy()
        letter_df = out[out["lane_treatment"] == "Letter"].copy()

        st.markdown("### Excel comparison workbook")

        def build_comparison_workbook():
            buf = io.BytesIO()
            with pd.ExcelWriter(buf, engine="openpyxl") as writer:
                out.to_excel(writer, sheet_name="All_NonPrivateFleet", index=False)
                letter_df.to_excel(writer, sheet_name="Letter_Lanes", index=False)
                rfp_df.to_excel(writer, sheet_name="RFP_Lanes", index=False)
                gpf_export.to_excel(writer, sheet_name="Private_Fleet", index=False)
                summary_df.to_excel(writer, sheet_name="Summary", index=False)
                excluded_summary_df.to_excel(writer, sheet_name="Excluded_Summary", index=False)
                excluded_detail_df.to_excel(writer, sheet_name="Excluded_Detail", index=False)
            buf.seek(0)
            return buf.getvalue()

        if st.button("Prepare Excel comparison workbook"):
            st.session_state["comparison_xlsx"] = build_comparison_workbook()

        if "comparison_xlsx" in st.session_state:
            st.download_button(
                label="⬇️ Download Excel workbook",
                data=st.session_state["comparison_xlsx"],
                file_name="lane_comparison.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        # =========================
        # Exports: RFP Template & Letters
        # =========================
        st.markdown("---")
        st.markdown("#### 📝 RFP Template & Letters")
        st.markdown(
            "Use this section when you are planning an **RFP event**.\n\n"
            "The **RFP Template (Excel)** contains:\n"
            "- A cover sheet with overview, instructions, and contact info\n"
            "- A bid timeline\n"
            "- A tab of RFP lanes with **Market Rate** and fields for carrier input\n"
            "- A location list tab\n\n"
            "The **RFP letters (by carrier)**:\n"
            "- Reference the total number of lanes in the RFP\n"
            "- Highlight the lanes that carrier currently services above benchmark\n"
            "- Show the **top 10% lanes with the largest variance**, plus summaries\n"
        )
        
        # --- Recompute lane treatments inside exports to ensure variables exist here ---
        out = out.copy()
        out = ensure_origin_dest(out)
        
        # If you stored overrides in session_state, use them; otherwise use a local variable you set in Configure tab.
        override_letter_lanes = st.session_state.get("override_letter_lanes", set())
        
        out["lane_key_norm"] = out["_lane"].astype(str).str.strip().str.upper()
        
        def classify_row(row):
            if row.get("action") != "NEGOTIATE":
                return "None"
            if pd.isna(row.get("benchmark_cost")) or float(row.get("benchmark_cost", 0) or 0) <= 0:
                return "None"
            if row["lane_key_norm"] in override_letter_lanes:
                return "Letter"
            return "RFP"
        
        out["lane_treatment"] = out.apply(classify_row, axis=1)
        
        rfp_df = out[out["lane_treatment"] == "RFP"].copy()
        letter_df = out[out["lane_treatment"] == "Letter"].copy()
        
        # =========================
        # Build RFP Template (Excel)
        # =========================
        st.subheader("📦 Build RFP Template (Excel)")
        
        overview_text = st.text_area(
            "RFP overview text (Tab 1): Please **revise the text below** to be specific to your company. Any edits you make here will be reflected in the RFP template.",
            value=(
                "Greif, Inc. is a global manufacturer of industrial packaging products "
                "and services, supplying fiber, steel, and plastic drums; intermediate "
                "bulk containers; corrugated and paper-based packaging; and containerboard "
                "to customers worldwide. The company also offers packaging accessories, "
                "transit protection products, and a range of reconditioning, recycling, "
                "and supply-chain services that support Circular Economy and sustainability initiatives. "
                "Greif serves major industries including chemicals, pharmaceuticals, agriculture, "
                "food and beverage, petroleum, and general manufacturing. With a broad production "
                "footprint and a strong commitment to operational excellence, product quality, and "
                "environmental stewardship, Greif is a trusted partner for safe, reliable, and "
                "efficient packaging solutions. Founded in 1877, the company is headquartered in Delaware, Ohio."
            ),
            help="This text will populate the Overview tab in the RFP."
        )
        
        # Use a reply-by field for the template (your instructions_text includes {reply_by})
        rfp_template_reply_by = st.text_input(
            "RFP template reply-by date (used in Instructions text below)",
            value="",
            help="Inserted into the Instructions section wherever {reply_by} is used."
        )
        
        instructions_text = st.text_area(
            "RFP Instructions (Tab 1): Please **revise the text below** to be specific to your company. Any edits you make here will be reflected in the RFP template.",
            value=(
                "Greif is continously looking to improve our ways of working and unlock efficiencies across our Supply Chain. "
                "Therefore, we are leveraging our volume and knowledge of market rates to complete a bid for several of our lanes and "
                "request that your company participate.\n\n"
                "In this document you will find three tabs:\n"
                "(1) Bid timeline - expected due dates and responses dates\n"
                "(2) All lanes that are open to bid - please enter bid cost as base rate charge and fuel surcharge as percentage of base rate\n"
                "(3) Locations covered by lanes included in this RFP - for reference only, no action required\n\n"
                "We request you complete and return this RFP by no later than {reply_by}."
            ),
            height=200
        )
        
        timeline_text = st.text_area(
            "Bid timeline (Tab 2): One line per milestone, format: Date, Milestone",
            value=(
                "4/22, Bid Release\n"
                "5/1, Round 1 Carrier Offers Due\n"
                "5/15, Round 1 [Company Name] Feedback to Carriers\n"
                "5/29, Round 2 Carrier Offers Due\n"
                "6/7, Round 2 [Company Name] Feedback to Carriers & Final Negotiations\n"
                "6/13, Final Awards"
            ),
            help="Each line should have a date and a milestone, separated by a comma."
        )
        
        rfp_loc_file = st.file_uploader(
            "Upload location list for RFP (optional – becomes last tab)",
            type=["xlsx", "xls", "csv"],
            key="rfp_locs"
        )
        
        # --- Timeline DF ---
        timeline_rows = []
        for raw in timeline_text.split("\n"):
            raw = raw.strip()
            if not raw:
                continue
            if "," in raw:
                date_str, milestone = raw.split(",", 1)
                timeline_rows.append({"Date": date_str.strip(), "Milestone": milestone.strip()})
            else:
                timeline_rows.append({"Date": raw, "Milestone": ""})
        timeline_df = pd.DataFrame(timeline_rows) if timeline_rows else pd.DataFrame(columns=["Date", "Milestone"])
        
        # --- Bid lanes export DF ---
        lane_cols = ["origin_city", "origin_state", "dest_city", "dest_state"]
        rfp_df = ensure_origin_dest(rfp_df)
        
        if not rfp_df.empty:
            rfp_group = (
                rfp_df.groupby(lane_cols, dropna=False)
                .agg(Shipment_Count=("_lane", "size"), Market_Rate=("benchmark_cost", "mean"))
                .reset_index()
            )
        
            bid_lanes_export = pd.DataFrame({
                "Zip ID": range(1, len(rfp_group) + 1),
                "Move Type": "US Domestic",
                "Origin City": rfp_group["origin_city"],
                "Origin State": rfp_group["origin_state"],
                "Origin Country": "US",
                "Destination City": rfp_group["dest_city"],
                "Destination State": rfp_group["dest_state"],
                "Destination Country": "US",
                "Shipment Count": rfp_group["Shipment_Count"],
                "Market Rate": rfp_group["Market_Rate"],
                "Flat rate (before fuel surcharge)": pd.NA,
                "Fuel surcharge (as percentage of flat rate)": pd.NA,
            })
        else:
            bid_lanes_export = pd.DataFrame(columns=[
                "Zip ID", "Move Type",
                "Origin City", "Origin State", "Origin Country",
                "Destination City", "Destination State", "Destination Country",
                "Shipment Count", "Market Rate",
                "Flat rate (before fuel surcharge)",
                "Fuel surcharge (as percentage of flat rate)",
            ])
        
        # --- Excel writer (your formatting preserved, but correctly indented/executable) ---
        from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
        
        if st.button("⬇️ Build RFP Template (Excel)", key="build_rfp_template"):
            # Location list
            if rfp_loc_file is not None:
                loc_df = read_any_cached(rfp_loc_file)
            else:
                loc_df = pd.DataFrame()
        
            overview_lines = [line for line in overview_text.split("\n") if line.strip()]
        
            # Fill in {reply_by} safely
            reply_by_text = rfp_template_reply_by.strip() if rfp_template_reply_by.strip() else "the requested RFP due date"
            instructions_filled = instructions_text.replace("{reply_by}", reply_by_text)
            instruction_lines = [line for line in instructions_filled.split("\n") if line.strip()]
        
            contact_lines = [
                "Name: _______________________________",
                "Title: ______________________________",
                "Company: ____________________________",
                "Email: _____________________________",
                "Phone: _____________________________",
            ]
        
            rfp_buf = io.BytesIO()
            with pd.ExcelWriter(rfp_buf, engine="openpyxl") as writer:
                # TAB 1
                pd.DataFrame().to_excel(writer, sheet_name="RFP Overview", index=False)
                ws = writer.sheets["RFP Overview"]
        
                thin_border = Border(
                    left=Side(style="thin", color="000000"),
                    right=Side(style="thin", color="000000"),
                    top=Side(style="thin", color="000000"),
                    bottom=Side(style="thin", color="000000"),
                )
        
                def apply_border(ws_, start_row, end_row, start_col="B", end_col="H"):
                    for row in ws_[f"{start_col}{start_row}:{end_col}{end_row}"]:
                        for cell in row:
                            cell.border = thin_border
        
                def add_text_section(ws_, title, lines, start_row, fill_color="EEECE1"):
                    title_cell = ws_[f"B{start_row}"]
                    title_cell.value = title
                    title_cell.font = Font(size=14, bold=True)
                    title_cell.alignment = Alignment(horizontal="left", vertical="bottom")
        
                    body_start = start_row + 1
                    body_height = max(5, len(lines) + 2)
                    body_end = body_start + body_height - 1
        
                    ws_.merge_cells(f"B{body_start}:H{body_end}")
                    cell = ws_[f"B{body_start}"]
                    cell.value = "\n".join(lines) if lines else ""
                    cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
                    cell.font = Font(size=12)
                    cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
        
                    apply_border(ws_, body_start, body_end)
                    return body_end + 2
        
                ws.column_dimensions["A"].width = 2
                ws.column_dimensions["B"].width = 60
                for col in ["C", "D", "E", "F", "G", "H"]:
                    ws.column_dimensions[col].width = 20
        
                current_row = 2
                current_row = add_text_section(ws, "1. Overview", overview_lines, current_row, fill_color="FFF2CC")
                current_row = add_text_section(ws, "2. Instructions", instruction_lines, current_row, fill_color="E2EFDA")
                current_row = add_text_section(ws, "3. Contact Information", contact_lines, current_row, fill_color="D9E1F2")
        
                # TAB 2: Timeline
                sheet_name_timeline = "Bid Timeline"
                timeline_df.to_excel(writer, sheet_name=sheet_name_timeline, index=False, startrow=3, startcol=1)
                ws_tl = writer.sheets[sheet_name_timeline]
        
                ws_tl["B2"].value = "Bid Timeline"
                ws_tl["B2"].font = Font(size=14, bold=True)
                ws_tl["B2"].alignment = Alignment(horizontal="left", vertical="bottom")
        
                start_row = 4
                end_row = start_row + len(timeline_df)
                start_col = "B"
                end_col = chr(ord(start_col) + len(timeline_df.columns) - 1) if len(timeline_df.columns) else "B"
        
                header_row = ws_tl[start_row]
                for cell in header_row:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    cell.fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
        
                for row in ws_tl.iter_rows(min_row=start_row + 1, max_row=end_row):
                    for cell in row:
                        cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
        
                for row in ws_tl[f"{start_col}{start_row}:{end_col}{end_row}"]:
                    for cell in row:
                        cell.border = thin_border
                        if cell.row != start_row:
                            cell.fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
        
                for col_cells in ws_tl.iter_cols(min_row=start_row, max_row=end_row, min_col=2, max_col=1 + len(timeline_df.columns)):
                    max_length = 0
                    col_letter = col_cells[0].column_letter
                    for cell in col_cells:
                        val = "" if cell.value is None else str(cell.value)
                        max_length = max(max_length, len(val))
                    ws_tl.column_dimensions[col_letter].width = max_length + 4
        
                for cell in ws_tl["A:A"]:
                    cell.border = Border()
                    cell.fill = PatternFill(fill_type=None)
                    cell.font = Font(size=11)
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                ws_tl.column_dimensions["A"].width = 2
        
                # TAB 3: Bid Lanes
                sheet_name_lanes = "Bid Lanes"
                bid_lanes_export.to_excel(writer, sheet_name=sheet_name_lanes, index=False, startrow=3, startcol=1)
                ws_bl = writer.sheets[sheet_name_lanes]
        
                ws_bl["B2"].value = "Bid Lanes"
                ws_bl["B2"].font = Font(size=14, bold=True)
                ws_bl["B2"].alignment = Alignment(horizontal="left", vertical="bottom")
        
                start_row_bl = 4
                end_row_bl = start_row_bl + len(bid_lanes_export)
                start_col_idx = 2
                num_cols = len(bid_lanes_export.columns)
                end_col_idx = start_col_idx + num_cols - 1
        
                header_cells = ws_bl[start_row_bl]
                col_idx_to_name = {cell.column: str(cell.value) for cell in header_cells}
        
                carrier_fill_cols = {
                    "Flat rate (before fuel surcharge)",
                    "Fuel surcharge (as percentage of flat rate)",
                }
        
                header_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
                data_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
                carrier_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        
                for cell in header_cells:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    cell.fill = header_fill
                    cell.border = thin_border
        
                for row in ws_bl.iter_rows(min_row=start_row_bl + 1, max_row=end_row_bl, min_col=start_col_idx, max_col=end_col_idx):
                    for cell in row:
                        header_name = col_idx_to_name.get(cell.column, "")
                        cell.fill = carrier_fill if header_name in carrier_fill_cols else data_fill
                        cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
                        cell.border = thin_border
        
                for col_idx in range(start_col_idx, end_col_idx + 1):
                    col_letter = ws_bl.cell(row=start_row_bl, column=col_idx).column_letter
                    max_length = 0
                    for r in range(start_row_bl, end_row_bl + 1):
                        v = ws_bl.cell(row=r, column=col_idx).value
                        max_length = max(max_length, len("" if v is None else str(v)))
                    ws_bl.column_dimensions[col_letter].width = max(15, max_length + 4)
        
                for cell in ws_bl["A:A"]:
                    cell.border = Border()
                    cell.fill = PatternFill(fill_type=None)
                    cell.font = Font(size=11)
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                ws_bl.column_dimensions["A"].width = 2
        
                # TAB 4: Location List
                if not loc_df.empty:
                    loc_df.to_excel(writer, sheet_name="Company Location List", index=False)
                else:
                    pd.DataFrame(columns=[
                        "Location Name", "Street Address", "City", "State", "Zip",
                        "Location Type", "Ledger Flag", "Platform"
                    ]).to_excel(writer, sheet_name="Company Location List", index=False)
        
            rfp_buf.seek(0)
            st.download_button(
                label="📥 Download RFP Template (Excel)",
                data=rfp_buf,
                file_name="RFP_Template.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        
        # =========================
        # RFP Letters (ZIP)
        # =========================
        st.markdown("---")
        st.subheader("✉️ Generate RFP Letters (by Carrier)")
        st.markdown(
            "This section will:\n"
            "- Use lanes marked as **RFP** in the comparison logic\n"
            "- Group lanes by carrier\n"
            "- Show unique lanes, frequency, company vs benchmark cost, and % over benchmark\n"
        )
        
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            rfp_sender_company = st.text_input("Your company name", value="Greif", key="rfp_sender_company")
        with col_b:
            rfp_sender_name = st.text_input("Your name", value="Your Name", key="rfp_sender_name")
        with col_c:
            rfp_sender_title = st.text_input("Your title", value="Procurement Manager", key="rfp_sender_title")
        
        rfp_reply_by = st.text_input("RFP reply-by date", value="", key="rfp_reply_by")
        
        rfp_default_body = (
            "Good afternoon,\n\n"
            "We are reaching out as we are completing an RFP for our Freight spend. "
            "The RFP that has been shared with you has a total {total_rfp_lanes} lanes. "
            "We appreciate your partnership and would like to continue leveraging your company as a strategic carrier.\n\n"
            "Of the lanes that are included in this RFP, please note there are {carrier_lane_count} "
            "lanes you service that are currently above benchmark rates at an average of {avg_over_pct} above benchmark. "
            "Our goal is to close the gap to benchmark rates as much as possible and would appreciate your careful consideration "
            "as you complete this RFP.\n\n"
            "For your reference and ease, the top 10% lanes with the greatest delta to benchmark are listed below with a summary of remaining lanes. "
            "If you would like further detail on the remaining lanes, please reach out.\n\n"
            "Please submit your responses by {reply_by}."
        )
        
        rfp_body_template = st.text_area(
            "RFP letter body (editable)",
            value=rfp_default_body,
            height=220,
            key="rfp_body_template"
        )
        
        include_privfleet_in_rfp_letters = st.checkbox("Include PRIVATE FLEET in RFP letters", value=False, key="include_pf_rfp_letters")
        
        if st.button("Build RFP letters (ZIP)", key="build_rfp_letters"):
            if rfp_df.empty:
                st.warning("No RFP lanes available to create letters.")
            else:
                zip_bytes = build_rfp_letters_zip(
                    rfp_df=rfp_df,
                    sender_company=rfp_sender_company,
                    sender_name=rfp_sender_name,
                    sender_title=rfp_sender_title,
                    reply_by=rfp_reply_by if rfp_reply_by else "the requested RFP due date",
                    rfp_body_template=rfp_body_template,
                    include_privfleet=include_privfleet_in_rfp_letters,
                )
                st.download_button(
                    label="⬇️ Download RFP letters (ZIP)",
                    data=zip_bytes,
                    file_name="rfp_letters.zip",
                    mime="application/zip"
                )
        
        # =========================
        # Negotiation Letters (ZIP)
        # =========================
        st.markdown("---")
        st.subheader("✉️ Generate Negotiation Letters")
        
        neg_letter_df = letter_df[letter_df["action"] == "NEGOTIATE"].copy()
        if neg_letter_df.empty:
            st.info(
                "No lanes are currently flagged for negotiation letters. "
                "Under the current rules and overrides, all lanes above benchmark are either routed to the RFP or excluded."
            )
        else:
            st.success(
                f"{neg_letter_df['carrier_name'].nunique()} carriers and {neg_letter_df['_lane'].nunique()} unique lanes are in scope."
            )
        
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            sender_company = st.text_input("Your company name", value="Greif", key="neg_sender_company")
        with col_b:
            sender_name = st.text_input("Your name", value="Your Name", key="neg_sender_name")
        with col_c:
            sender_title = st.text_input("Your title", value="Procurement Manager", key="neg_sender_title")
        
        col_d, col_e = st.columns(2)
        with col_d:
            reply_by = st.text_input("Reply-by date", value="", key="neg_reply_by")
        with col_e:
            include_privfleet_in_letters = st.checkbox("Include PRIVATE FLEET in letters", value=False, key="include_pf_neg_letters")
        
        default_intro = (
            "We appreciate your service and would like to continue our partnership. "
            "We have identified {num_lanes} lanes where the rates we are charged are "
            "above market by an average of {avg_over_pct} and would like to review rates. "
            "Please review and provide your best offer by {reply_by}."
        )
        
        st.markdown("**Negotiation letter body**")
        st.caption("Placeholders: {num_lanes}, {avg_over_pct}, {reply_by}")
        
        letter_body_template = st.text_area(
            "Negotiation letter body (editable)",
            value=default_intro,
            height=120,
            key="neg_letter_body_template",
        )
        
        if st.button("Build negotiation letters (ZIP)", key="build_neg_letters"):
            combined_for_letters = neg_letter_df.assign(source="NON_GREIF")
        
            if include_privfleet_in_letters and isinstance(gpf_export, pd.DataFrame) and (not gpf_export.empty):
                greif_neg = gpf_export[gpf_export["action"] == "NEGOTIATE"].copy()
                if not greif_neg.empty:
                    greif_neg["source"] = "GREIF"
                    combined_for_letters = pd.concat([combined_for_letters, greif_neg], ignore_index=True, sort=False)
        
            if combined_for_letters.empty:
                st.warning("No negotiation letters were generated because no lanes are currently classified as Letter + NEGOTIATE.")
            else:
                zip_bytes = build_letters_zip(
                    df_all=combined_for_letters,
                    include_privfleet=include_privfleet_in_letters,
                    sender_company=sender_company,
                    sender_name=sender_name,
                    sender_title=sender_title,
                    reply_by=reply_by if reply_by else "7 business days from receipt",
                    body_template=letter_body_template,
                )
                st.download_button(
                    label="⬇️ Download negotiation letters (ZIP)",
                    data=zip_bytes,
                    file_name="negotiation_letters.zip",
                    mime="application/zip"
                )
    st.button("← Back to Results (click tab above)", key="back_exports")
